from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from PIL import Image
from docx import Document
from docx.oxml import parse_xml
from reportlab.pdfgen import canvas

from pdf2word_engine.document_checks import (
    DocumentCheckError,
    assert_rendered_page_count,
    assert_source_first_docx_contract,
    inspect_docx_structure,
)
from pdf2word_engine.models import PageBlock, PageModel, PageSize, PdfKind
from pdf2word_engine.word import create_positioned_editable_docx
from pdf2word_engine.word_render import render_docx_with_microsoft_word


def _source_first_model(image_path: Path | None = None) -> PageModel:
    blocks = [
        PageBlock(
            "paragraph",
            "editable_paragraph",
            (60, 90, 450, 150),
            1,
            1,
            text="这是可编辑的原生 Word 正文。",
            style={"semantic_role": "question_body", "font_size_pt": 10.5},
        )
    ]
    if image_path is not None:
        blocks.append(
            PageBlock(
                "formula",
                "formula_image",
                (180, 180, 320, 230),
                2,
                2,
                asset_path=str(image_path),
                style={"semantic_role": "formula", "fallback_method": "source_crop"},
            )
        )
    return PageModel(
        schema_version=4,
        page_index=0,
        size=PageSize(516, 729),
        source_type=PdfKind.OUTLINED,
        source_image_width_px=1032,
        source_image_height_px=1458,
        page_class="ordinary_question",
        reconstruction_mode="native_word_paragraphs_with_clean_source_region_fallbacks",
        blocks=blocks,
    )


class DocumentChecksTests(unittest.TestCase):
    def test_source_first_structure_reports_editable_paragraphs_and_fallbacks(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            image = root / "formula.png"
            Image.new("RGB", (140, 50), "white").save(image)
            docx = create_positioned_editable_docx([_source_first_model(image)], root / "candidate.docx")
            report = assert_source_first_docx_contract(docx, minimum_editable_characters=10)

        self.assertGreaterEqual(report.native_frame_paragraphs, 1)
        self.assertEqual(report.legacy_vml_text_boxes, 0)
        self.assertEqual(report.fallback_images, 1)

    def test_old_line_text_box_layout_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            docx = Path(temp) / "legacy.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run()._r.append(
                parse_xml(
                    '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                    'xmlns:v="urn:schemas-microsoft-com:vml"><v:shape id="legacy" type="#_x0000_t202">'
                    '<v:textbox><w:txbxContent><w:p><w:r><w:t>旧文本框</w:t></w:r></w:p>'
                    '</w:txbxContent></v:textbox></v:shape></w:pict>'
                )
            )
            document.save(docx)
            report = inspect_docx_structure(docx)
            with self.assertRaisesRegex(DocumentCheckError, "旧式 VML 文字框"):
                assert_source_first_docx_contract(docx)

        self.assertGreater(report.legacy_vml_text_boxes, 0)

    def test_structure_report_counts_native_math_and_inline_decoration(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            image = root / "talk.png"
            Image.new("RGBA", (100, 50), (239, 22, 139, 255)).save(image)
            model = PageModel(
                schema_version=8,
                page_index=0,
                size=PageSize(516, 729),
                source_type=PdfKind.OUTLINED,
                source_image_width_px=1032,
                source_image_height_px=1458,
                page_class="ordinary_question",
                blocks=[
                    PageBlock(
                        "talk",
                        "talk_label_image",
                        (120, 180, 220, 230),
                        0,
                        0,
                        text="解析",
                        style={"inline_decorative": True, "inline_host_block_id": "body"},
                        asset_path=str(image),
                        fallback_mode="talk_label_source_image",
                    ),
                    PageBlock(
                        "body",
                        "editable_callout_body",
                        (80, 170, 900, 300),
                        0,
                        1,
                        text="可编辑分数2/5。",
                        style={"contains_inline_label": True, "first_line_indent_px": 40.0},
                    ),
                ],
            )
            docx = create_positioned_editable_docx([model], root / "native-math.docx")
            report = inspect_docx_structure(docx)

        self.assertEqual(report.native_math_fractions, 1)
        self.assertEqual(report.inline_decorative_images, 1)
        self.assertEqual(report.fallback_images, 1)

    def test_page_count_gate_rejects_unexpected_overflow(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            pdf = Path(temp) / "two-pages.pdf"
            document = canvas.Canvas(str(pdf))
            document.drawString(72, 720, "one")
            document.showPage()
            document.drawString(72, 720, "two")
            document.save()
            with self.assertRaisesRegex(DocumentCheckError, "分页门禁失败"):
                assert_rendered_page_count(pdf, 1)

    def test_word_renderer_reports_missing_word_dependency_or_document(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            with self.assertRaisesRegex(DocumentCheckError, "Microsoft Word"):
                render_docx_with_microsoft_word(Path(temp) / "missing.docx", Path(temp) / "word.pdf")
