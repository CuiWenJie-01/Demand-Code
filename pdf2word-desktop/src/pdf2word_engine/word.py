"""DOCX generation helpers for editable Word output."""

from __future__ import annotations

from math import ceil
from pathlib import Path
import re
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Emu, Pt
from pypdf import PdfReader

from .errors import OcrRequiredError
from .conflicts import resolve_page_model_conflicts
from .models import PageBlock, PageModel, PageSize, PdfKind


EMU_PER_POINT = 12_700


def _set_page_geometry(section: object, size: PageSize) -> None:
    section.page_width = Emu(round(size.width_pt * EMU_PER_POINT))  # type: ignore[attr-defined]
    section.page_height = Emu(round(size.height_pt * EMU_PER_POINT))  # type: ignore[attr-defined]
    section.top_margin = Emu(0)  # type: ignore[attr-defined]
    section.bottom_margin = Emu(0)  # type: ignore[attr-defined]
    section.left_margin = Emu(0)  # type: ignore[attr-defined]
    section.right_margin = Emu(0)  # type: ignore[attr-defined]
    section.header_distance = Emu(0)  # type: ignore[attr-defined]
    section.footer_distance = Emu(0)  # type: ignore[attr-defined]


def _format_page_paragraph(paragraph: object) -> None:
    paragraph.paragraph_format.space_before = Pt(0)  # type: ignore[attr-defined]
    paragraph.paragraph_format.space_after = Pt(0)  # type: ignore[attr-defined]
    paragraph.paragraph_format.line_spacing = 1  # type: ignore[attr-defined]


def _page_coordinate_scale(model: PageModel) -> tuple[float, float]:
    width = model.source_image_width_px or model.size.width_pt
    height = model.source_image_height_px or model.size.height_pt
    return model.size.width_pt / width, model.size.height_pt / height


def _shape_style(block: PageBlock, model: PageModel) -> tuple[str, float, str, str]:
    scale_x, scale_y = _page_coordinate_scale(model)
    left, top, right, bottom = block.bbox
    width = max(1.0, (right - left) * scale_x)
    height = max(1.0, (bottom - top) * scale_y)
    try:
        configured_width_padding_pt = float(block.style.get("textbox_width_padding_pt"))
        if configured_width_padding_pt > 0:
            width += configured_width_padding_pt
    except (TypeError, ValueError):
        pass
    # OCR regions can contain several visual lines.  Deriving the font straight
    # from total region height makes multi-line blocks enormous, so estimate the
    # line count from the available width and content length first.
    text_length = len("".join((block.text or "").split()))
    estimated_char_width = 8.5
    estimated_chars_per_line = max(1, int(width / estimated_char_width))
    estimated_lines = max(1, ceil(text_length / estimated_chars_per_line))
    font_pt = max(7.0, min(12.0, (height / estimated_lines) * 0.65))
    block_type = block.block_type.lower()
    semantic_role = str(block.style.get("semantic_role", ""))
    # PP-Structure's line boxes are commonly only 8–11 pt high after page
    # scaling.  LibreOffice (and some Word versions) clips CJK glyphs in a VML
    # textbox this shallow.  A line is already the smallest editable unit, so
    # keep it to one line with a compact, but legible, font and give its shape
    # enough vertical room for the font metrics.
    if block_type == "text_line":
        font_pt = min(9.0, max(8.0, height * 0.72))
        height = max(height, font_pt * 1.55)
    color = "000000"
    fill = "f"
    if block_type == "paragraph_title":
        color = "EF168B"
        font_pt = max(font_pt, 8.0)
    elif block_type == "header":
        color = "000000"
        font_pt = max(font_pt, 8.5)
    elif block_type in {"aside_text", "number"}:
        color = "555555"
        font_pt = min(font_pt, 8.5)
    try:
        configured_font_pt = float(block.style.get("font_size_pt"))
        if configured_font_pt > 0:
            font_pt = configured_font_pt
    except (TypeError, ValueError):
        pass
    # One-line OCR boxes are intentionally positioned as one-line boxes.  Do
    # not let Word wrap a short answer into a neighbouring line.  Tighten a
    # little first; a page conflict check can choose a source-image fallback
    # for genuinely unfit content.
    if block_type == "text_line" and text_length:
        safe_font = width / max(1.0, text_length * 0.93)
        font_pt = min(font_pt, max(6.5, safe_font))
    try:
        configured_min_height_pt = float(block.style.get("textbox_min_height_pt"))
        if configured_min_height_pt > 0:
            height = max(height, configured_min_height_pt)
    except (TypeError, ValueError):
        pass
    configured_color = block.style.get("font_color")
    if isinstance(configured_color, str) and len(configured_color) == 6:
        color = configured_color.upper()
    z_index = 0 if block_type == "watermark" or block.style.get("render_behind_text") else max(1, block.z_index + 1)
    style = (
        "position:absolute;"
        "mso-position-horizontal-relative:page;"
        "mso-position-vertical-relative:page;"
        f"margin-left:{left * scale_x:.2f}pt;"
        f"margin-top:{top * scale_y:.2f}pt;"
        f"width:{width:.2f}pt;"
        f"height:{height:.2f}pt;"
        f"z-index:{z_index};"
    )
    return style, font_pt, color, fill


def _append_textbox(paragraph: object, model: PageModel, block: PageBlock) -> None:
    text = (block.text or "").replace("\r", " ").replace("\n", " ").strip()
    if not text:
        return
    style, font_pt, color, fill = _shape_style(block, model)
    line_height = font_pt * 1.1 if block.block_type.lower() in {"paragraph_title", "header"} else 12.0
    accent_length = block.style.get("accent_length")
    try:
        accent_length = max(0, min(len(text), int(accent_length))) if accent_length is not None else 0
    except (TypeError, ValueError):
        accent_length = 0
    bold_prefix_length = block.style.get("bold_prefix_length")
    try:
        bold_prefix_length = max(0, min(len(text), int(bold_prefix_length))) if bold_prefix_length is not None else 0
    except (TypeError, ValueError):
        bold_prefix_length = 0
    character_spacing_twips = block.style.get("character_spacing_twips")
    try:
        character_spacing_twips = int(character_spacing_twips) if character_spacing_twips is not None else 0
    except (TypeError, ValueError):
        character_spacing_twips = 0
    runs = _textbox_runs(
        text,
        accent_length=accent_length,
        bold_prefix_length=bold_prefix_length,
        font_pt=font_pt,
        color=color,
        character_spacing_twips=character_spacing_twips,
        east_asia_font=str(block.style.get("font_name_east_asia", "SimSun")),
        ascii_font=str(block.style.get("font_name_ascii", "Times New Roman")),
    )
    requested_alignment = str(block.style.get("text_alignment", "")).lower()
    if requested_alignment in {"left", "center", "right", "distribute"}:
        alignment = f'<w:jc w:val="{requested_alignment}"/>'
    else:
        alignment = '<w:jc w:val="distribute"/>' if _should_distribute(block, model, text) else '<w:jc w:val="left"/>'
    xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml">
      <v:shape id="pdf2word_text_{escape(block.block_id)}" type="#_x0000_t202" style="{style}" stroked="f" filled="{fill}">
        <v:textbox inset="0pt,0pt,0pt,0pt" style="mso-fit-shape-to-text:f"><w:txbxContent><w:p><w:pPr><w:spacing w:before="0" w:after="0" w:line="{round(line_height * 20)}"/>{alignment}</w:pPr>
          {runs}
        </w:p></w:txbxContent></v:textbox>
      </v:shape>
    </w:pict>'''
    paragraph.add_run()._r.append(parse_xml(xml))


def _should_distribute(block: PageBlock, model: PageModel, text: str) -> bool:
    """Permit spacing expansion only for near-full editable prose lines."""

    if not block.style.get("justify_to_bbox") or len("".join(text.split())) < 18:
        return False
    role = str(block.style.get("semantic_role", ""))
    if role in {"callout_label", "callout_answer", "answer_blank", "solution_short_body", "callout_body_fragment"}:
        return False
    page_width = model.source_image_width_px or model.size.width_pt
    return (block.bbox[2] - block.bbox[0]) >= page_width * 0.62


def _textbox_runs(
    text: str,
    *,
    accent_length: int,
    bold_prefix_length: int,
    font_pt: float,
    color: str,
    character_spacing_twips: int = 0,
    east_asia_font: str = "SimSun",
    ascii_font: str = "Times New Roman",
) -> str:
    """Serialize editable runs with independently coloured and bold prefixes."""

    def run(value: str, run_color: str, *, bold: bool) -> str:
        # Be explicit for every non-title run: some Word/LibreOffice versions
        # otherwise inherit the preceding run's bold state inside a VML box.
        bold_xml = "<w:b/>" if bold else '<w:b w:val="0"/>'
        character_spacing_xml = f'<w:spacing w:val="{character_spacing_twips}"/>' if character_spacing_twips else ""
        return (
            f'<w:r><w:rPr><w:rFonts w:ascii="{escape(ascii_font)}" w:hAnsi="{escape(ascii_font)}" w:eastAsia="{escape(east_asia_font)}"/>'
            f'{bold_xml}{character_spacing_xml}<w:color w:val="{run_color}"/><w:sz w:val="{round(font_pt * 2)}"/></w:rPr>'
            f'<w:t xml:space="preserve">{escape(value)}</w:t></w:r>'
        )

    boundaries = sorted({0, len(text), accent_length, bold_prefix_length})
    return "".join(
        run(
            text[start:end],
            "EF168B" if start < accent_length else color,
            bold=start < bold_prefix_length,
        )
        for start, end in zip(boundaries, boundaries[1:])
        if start < end
    )


def _append_toc_entry(paragraph: object, model: PageModel, block: PageBlock) -> None:
    """Write one editable TOC-style paragraph with a real dot-leader tab."""

    style, font_pt, color, fill = _shape_style(block, model)
    scale_x, _ = _page_coordinate_scale(model)
    width_pt = max(1.0, (block.bbox[2] - block.bbox[0]) * scale_x)
    tab_position = max(120, round((width_pt - 3.0) * 20))
    chapter = escape(str(block.style.get("toc_chapter", "")))
    title = escape(str(block.style.get("toc_title", "")))
    page_number = escape(str(block.style.get("toc_page", "")))
    bookmark = str(block.style.get("target_bookmark", "")).strip()

    def run(value: str, *, ascii_font: str = "STSong") -> str:
        return (
            '<w:r><w:rPr>'
            f'<w:rFonts w:ascii="{escape(ascii_font)}" w:hAnsi="{escape(ascii_font)}" w:eastAsia="STSong"/>'
            f'<w:color w:val="{color}"/><w:sz w:val="{round(font_pt * 2)}"/>'
            '</w:rPr>'
            f'<w:t xml:space="preserve">{value}</w:t></w:r>'
        )

    content = run(f"{chapter}　{title}") + run("<TAB>").replace('<w:t xml:space="preserve">&lt;TAB&gt;</w:t>', '<w:tab/>') + run(page_number, ascii_font="Arial")
    if bookmark:
        content = f'<w:hyperlink w:anchor="{escape(bookmark)}" w:history="1">{content}</w:hyperlink>'
    xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml">
      <v:shape id="pdf2word_toc_{escape(block.block_id)}" type="#_x0000_t202" style="{style}" stroked="f" filled="{fill}">
        <v:textbox inset="0pt,0pt,0pt,0pt" style="mso-fit-shape-to-text:f"><w:txbxContent><w:p><w:pPr>
          <w:pStyle w:val="SourceTOC1"/><w:spacing w:before="0" w:after="0" w:line="260"/>
          <w:tabs><w:tab w:val="right" w:leader="dot" w:pos="{tab_position}"/></w:tabs><w:jc w:val="left"/>
        </w:pPr>{content}</w:p></w:txbxContent></v:textbox>
      </v:shape>
    </w:pict>'''
    paragraph.add_run()._r.append(parse_xml(xml))


def _toc_native_run_xml(value: str, *, font_pt: float, color: str, ascii_font: str = "STSong") -> str:
    return (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr>'
        f'<w:rFonts w:ascii="{escape(ascii_font)}" w:hAnsi="{escape(ascii_font)}" w:eastAsia="STSong"/>'
        f'<w:color w:val="{escape(color)}"/><w:sz w:val="{round(font_pt * 2)}"/>'
        '</w:rPr>'
        f'<w:t xml:space="preserve">{escape(value)}</w:t></w:r>'
    )


def _append_native_toc_page(document: Document, model: PageModel) -> None:
    """Render TOC text as normal Word paragraphs, not floating text boxes."""

    scale_x, scale_y = _page_coordinate_scale(model)
    decorations = [block for block in model.blocks if block.asset_path and block.block_type == "image"]
    text_blocks = sorted(
        (block for block in model.blocks if block.block_type in {"toc_group", "toc_entry"}),
        key=lambda item: (item.bbox[1], item.reading_order),
    )
    if not text_blocks:
        page_canvas = document.add_paragraph()
        _format_page_paragraph(page_canvas)
        for decoration in decorations:
            _append_fallback_image(page_canvas, model, decoration)
        return

    first_top = max(1.0, text_blocks[0].bbox[1] * scale_y)
    canvas_height = 1.0
    decoration_canvas = document.add_paragraph()
    _format_page_paragraph(decoration_canvas)
    decoration_canvas.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    decoration_canvas.paragraph_format.line_spacing = Pt(canvas_height)
    decoration_canvas.paragraph_format.space_before = Pt(0)
    decoration_canvas.paragraph_format.space_after = Pt(0)
    for decoration in decorations:
        _append_fallback_image(decoration_canvas, model, decoration)
        bookmark = str(decoration.style.get("bookmark_name", "")).strip()
        if bookmark:
            _append_bookmark(decoration_canvas, bookmark, model.page_index * 10 + 1)

    # Use paragraph spacing for the large top offset.  LibreOffice caps an
    # oversized exact line-height spacer, which previously moved the first TOC
    # group underneath the decoration strip and made it appear missing.
    cursor = canvas_height
    for block in text_blocks:
        top = block.bbox[1] * scale_y
        height = max(float(block.style.get("textbox_min_height_pt", 14.0)), (block.bbox[3] - block.bbox[1]) * scale_y)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(max(0.0, top - cursor))
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(height)
        paragraph.paragraph_format.keep_together = True
        left = block.bbox[0] * scale_x
        right = block.bbox[2] * scale_x
        paragraph.paragraph_format.left_indent = Pt(left)
        paragraph.paragraph_format.right_indent = Pt(max(0.0, model.size.width_pt - right))
        font_pt = float(block.style.get("font_size_pt", 10.5))
        color = str(block.style.get("font_color", "F4008A"))
        if block.block_type == "toc_group":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            wrapper = parse_xml(
                '<w:root xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                + _toc_native_run_xml(block.text or "", font_pt=font_pt, color=color)
                + '</w:root>'
            )
            for child in list(wrapper):
                paragraph._p.append(child)
        else:
            paragraph.style = document.styles["Source TOC 1"]
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Pt(max(left + 20.0, right - 3.0)),
                WD_TAB_ALIGNMENT.RIGHT,
                WD_TAB_LEADER.DOTS,
            )
            chapter = str(block.style.get("toc_chapter", ""))
            title = str(block.style.get("toc_title", ""))
            page_number = str(block.style.get("toc_page", ""))
            content = (
                _toc_native_run_xml(f"{chapter}　{title}", font_pt=font_pt, color=color)
                + _toc_native_run_xml("<TAB>", font_pt=font_pt, color=color).replace(
                    '<w:t xml:space="preserve">&lt;TAB&gt;</w:t>', '<w:tab/>'
                )
                + _toc_native_run_xml(page_number, font_pt=font_pt, color=color, ascii_font="Arial")
            )
            bookmark = str(block.style.get("target_bookmark", "")).strip()
            if bookmark:
                element = parse_xml(
                    f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                    f'w:anchor="{escape(bookmark)}" w:history="1">{content}</w:hyperlink>'
                )
                paragraph._p.append(element)
            else:
                wrapper = parse_xml(
                    '<w:root xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    + content
                    + '</w:root>'
                )
                for child in list(wrapper):
                    paragraph._p.append(child)
        cursor = top + height


def _append_bookmark(paragraph: object, name: str, bookmark_id: int) -> None:
    """Add a zero-layout internal-link target to the current source page."""

    safe_name = re.sub(r"[^A-Za-z0-9_]", "_", name)[:40]
    # Adjacent start/end elements are a valid zero-layout target.  A previous
    # hidden zero-width run caused LibreOffice to drop VML image shapes sharing
    # the paragraph, yielding blank cover pages during render verification.
    xml = f'''<w:bookmarkStart xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:id="{bookmark_id}" w:name="{escape(safe_name)}"/>
    <w:bookmarkEnd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:id="{bookmark_id}"/>'''
    wrapper = parse_xml(f'<w:root xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{xml}</w:root>')
    for child in list(wrapper):
        paragraph._p.append(child)  # type: ignore[attr-defined]


def _ensure_source_styles(document: Document) -> None:
    """Install explicit styles used by the source-faithful Word structure."""

    if "Source TOC 1" not in document.styles:
        toc = document.styles.add_style("Source TOC 1", WD_STYLE_TYPE.PARAGRAPH)
        toc._element.set(qn("w:styleId"), "SourceTOC1")
        toc.font.name = "STSong"
        toc._element.rPr.rFonts.set(qn("w:eastAsia"), "STSong")
        toc.font.size = Pt(10.5)
        toc.paragraph_format.space_before = Pt(0)
        toc.paragraph_format.space_after = Pt(0)
    if "Source Chapter Anchor" not in document.styles:
        anchor = document.styles.add_style("Source Chapter Anchor", WD_STYLE_TYPE.PARAGRAPH)
        anchor._element.set(qn("w:styleId"), "SourceChapterAnchor")
        anchor.base_style = document.styles["Heading 1"]
        anchor.font.hidden = True
        anchor.font.size = Pt(1)
        anchor.paragraph_format.space_before = Pt(0)
        anchor.paragraph_format.space_after = Pt(0)


def _append_fallback_image(paragraph: object, model: PageModel, block: PageBlock) -> None:
    if not block.asset_path or not Path(block.asset_path).is_file():
        return
    style, _, _, _ = _shape_style(block, model)
    relationship_id, _ = paragraph.part.get_or_add_image(str(block.asset_path))  # type: ignore[attr-defined]
    xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office"
        xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <v:shape id="pdf2word_image_{escape(block.block_id)}" type="#_x0000_t75" style="{style}" stroked="f">
        <v:imagedata r:id="{relationship_id}" o:title=""/>
      </v:shape>
    </w:pict>'''
    paragraph.add_run()._r.append(parse_xml(xml))


def _append_vertical_sidebar_text(paragraph: object, model: PageModel, block: PageBlock) -> None:
    """Write a vertical sidebar label as individually editable glyph boxes."""

    text = "".join((block.text or "").split())
    if not text:
        return
    left, top, right, bottom = block.bbox
    glyph_height = min(right - left, (bottom - top) / len(text))
    try:
        configured_height = float(block.style.get("glyph_height_px"))
        if configured_height > 0:
            glyph_height = min(configured_height, (bottom - top) / len(text))
    except (TypeError, ValueError):
        pass
    try:
        section_break_after = int(block.style.get("section_break_after"))
        section_break_after = max(0, min(len(text) - 1, section_break_after))
    except (TypeError, ValueError):
        section_break_after = 0
    try:
        section_gap = max(0.0, float(block.style.get("section_gap_px"))) if section_break_after else 0.0
    except (TypeError, ValueError):
        section_gap = 0.0
    usable_height = max(0.0, bottom - top - glyph_height - section_gap)
    step = 0.0 if len(text) == 1 else usable_height / (len(text) - 1)
    for index, glyph in enumerate(text):
        glyph_top = top + index * step + (section_gap if index >= section_break_after else 0.0)
        glyph_block = PageBlock(
            block_id=f"{block.block_id}-glyph-{index + 1}",
            block_type="text_line",
            bbox=(left, glyph_top, right, glyph_top + glyph_height),
            z_index=block.z_index + index,
            reading_order=block.reading_order + index,
            text=glyph,
            style=block.style,
        )
        _append_textbox(paragraph, model, glyph_block)


def _append_sidebar_accent_rule(paragraph: object, model: PageModel, block: PageBlock) -> None:
    """Render the page-number accent as an editable VML vector rectangle."""

    style, _, _, _ = _shape_style(block, model)
    fill_color = str(block.style.get("fill_color", "EF168B"))
    xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml">
      <v:rect id="pdf2word_sidebar_rule_{escape(block.block_id)}" style="{style}" stroked="f" fillcolor="#{escape(fill_color)}"/>
    </w:pict>'''
    paragraph.add_run()._r.append(parse_xml(xml))


def create_positioned_editable_docx(models: list[PageModel], output_path: str | Path) -> Path:
    """Build an editable DOCX using OCR coordinates and image fallbacks.

    Textual blocks become editable VML text boxes. Complex non-text blocks use
    their cropped image fallback so the document stays usable even when Word
    cannot reproduce a region structurally.
    """

    if not models:
        raise ValueError("没有可写入 DOCX 的 PageModel。")
    for model in models:
        resolve_page_model_conflicts(model)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _ensure_source_styles(document)
    for position, model in enumerate(sorted(models, key=lambda item: item.page_index)):
        section = document.sections[0] if position == 0 else document.add_section(WD_SECTION.NEW_PAGE)
        _set_page_geometry(section, model.size)
        if model.page_class == "table_of_contents":
            _append_native_toc_page(document, model)
            continue
        # All page objects are absolutely positioned. Keeping them in one host
        # paragraph prevents invisible flow paragraphs from accumulating and
        # pushing the final positioned objects onto an unwanted next page.
        page_canvas = document.add_paragraph()
        _format_page_paragraph(page_canvas)
        bookmark_names = [
            str(block.style.get("bookmark_name"))
            for block in model.blocks
            if block.style.get("bookmark_name")
        ]
        for bookmark_offset, bookmark_name in enumerate(dict.fromkeys(bookmark_names)):
            _append_bookmark(page_canvas, bookmark_name, position * 10 + bookmark_offset + 1)
        for block in sorted(model.blocks, key=lambda item: (item.z_index, item.reading_order)):
            if block.block_type == "toc_entry":
                _append_toc_entry(page_canvas, model, block)
                continue
            if block.block_type == "sidebar_vertical_text":
                _append_vertical_sidebar_text(page_canvas, model, block)
                continue
            if block.block_type == "sidebar_accent_rule":
                _append_sidebar_accent_rule(page_canvas, model, block)
                continue
            if block.asset_path:
                _append_fallback_image(page_canvas, model, block)
            if block.text and not block.asset_path and block.block_type.lower() not in {"image", "chart", "logo", "watermark"}:
                _append_textbox(page_canvas, model, block)
    document.save(output)
    return output


def _extract_page_text(source: str | Path) -> list[str]:
    reader = PdfReader(str(source), strict=False)
    if reader.is_encrypted:
        return []
    result: list[str] = []
    for page in reader.pages:
        try:
            result.append((page.extract_text(extraction_mode="layout") or "").strip())
        except TypeError:
            result.append((page.extract_text() or "").strip())
        except Exception:
            result.append("")
    return result


def create_basic_editable_docx(
    source: str | Path,
    *,
    page_sizes: list[PageSize],
    kind: PdfKind,
    page_indices: list[int],
    output_path: str | Path,
) -> Path:
    """Create an editable baseline for PDFs with an actual text layer.

    OCR-driven absolute-positioned layout is intentionally not emulated here.
    For outlined/scanned input callers receive an explicit error until the
    PaddleOCR PageModel adapter is installed and validated.
    """

    texts = _extract_page_text(source)
    selected = [texts[index] if index < len(texts) else "" for index in page_indices]
    if kind in {PdfKind.OUTLINED, PdfKind.SCANNED} or not any(selected):
        raise OcrRequiredError(
            "该 PDF 没有可靠文字层。需要安装并配置 PaddleOCR 后才能生成可编辑 Word。"
        )

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for position, (page_index, text) in enumerate(zip(page_indices, selected, strict=True)):
        if position == 0:
            section = document.sections[0]
            paragraph = document.add_paragraph()
        else:
            section = document.add_section(WD_SECTION.NEW_PAGE)
            paragraph = document.add_paragraph()
        _set_page_geometry(section, page_sizes[page_index])
        _format_page_paragraph(paragraph)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if lines:
            paragraph.add_run(lines[0])
            for line in lines[1:]:
                next_paragraph = document.add_paragraph()
                _format_page_paragraph(next_paragraph)
                next_paragraph.add_run(line)
        else:
            paragraph.add_run("[此页未提取到文字]")
    document.save(output)
    return output
